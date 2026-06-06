from pathlib import Path
import sys

sys.path.insert(0, r"C:\Users\C0SU\.cache\codex-runtimes\codex-primary-runtime\dependencies\python")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether

OUT = Path("output/module-one")
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Oscar_Vargas_Modulo_1_Organisational_Behaviour.docx"
PDF = OUT / "Oscar_Vargas_Modulo_1_Organisational_Behaviour.pdf"

INK = "151820"
LIME = "B9F227"
CYAN = "5FE7FF"
MUTED = "596273"
LIGHT = "F3F5F0"

sections = [
    ("1. Cómo usar esta guía", [
        "Esta guía acompaña el primer módulo de Organisational Behaviour CRKC 5001. Está diseñada para leer con calma, refrescar la memoria y convertir cada concepto en una observación aplicable.",
        "El recorrido recomendado es sencillo: leer una sección, cerrar el documento, explicar la idea con tus propias palabras y después aplicarla a una organización que conozcas. La memoria mejora cuando recuperamos activamente una idea, no cuando solamente la releemos."
    ]),
    ("2. Qué es el comportamiento organizacional", [
        "El comportamiento organizacional estudia cómo actúan las personas dentro de organizaciones y cómo esa conducta influye en resultados, cultura, bienestar y sostenibilidad. Su unidad de análisis no es solamente la persona. También observa equipos, estructuras, procesos y el contexto que condiciona las decisiones.",
        "La disciplina combina conocimientos de psicología, sociología, gestión y ciencias del comportamiento. Su valor práctico consiste en ayudar a describir lo que ocurre, explicar por qué ocurre y diseñar mejores condiciones para el trabajo."
    ]),
    ("3. Tres niveles para observar una organización", [
        "Nivel individual. Examina personalidad, valores, habilidades, percepción, motivación y aprendizaje. Ayuda a comprender por qué dos personas responden de manera distinta ante una situación similar.",
        "Nivel grupal. Examina normas, roles, cohesión, comunicación, poder, conflicto y coordinación. Permite entender por qué un conjunto de personas no siempre funciona como equipo.",
        "Nivel organizacional. Examina estructura, cultura, tecnología, procesos, estrategia y entorno. Explica cómo el sistema amplifica o limita la conducta de personas y equipos."
    ]),
    ("4. Naturaleza del trabajo y rol del manager", [
        "El trabajo no es solamente una lista de tareas. También produce identidad, relaciones, aprendizaje y sentido. El diseño del trabajo influye en autonomía, responsabilidad, retroalimentación, carga y bienestar.",
        "El manager contemporáneo coordina recursos y resultados, pero también diseña condiciones. Debe comunicar, decidir, desarrollar personas, resolver tensiones y ajustar el sistema. Gestionar tareas sin comprender conducta humana suele producir cumplimiento superficial, baja confianza o desgaste."
    ]),
    ("5. Cultura y clima organizacional", [
        "La cultura organizacional representa los supuestos, valores, normas y significados compartidos que explican cómo se hacen realmente las cosas. Es profunda, histórica y relativamente estable.",
        "El clima organizacional representa cómo las personas perciben su experiencia cotidiana: liderazgo, comunicación, reconocimiento, justicia, autonomía, carga y seguridad psicológica. Puede cambiar con mayor rapidez que la cultura.",
        "Una organización puede declarar valores positivos y, al mismo tiempo, mantener un clima negativo. Esa diferencia convierte al clima en una señal práctica para detectar cómo se vive la cultura."
    ]),
    ("6. Cómo medir el clima de forma profesional", [
        "Una encuesta de clima no debe comenzar escribiendo preguntas. Primero debe definir qué decisión informará, qué población participará, qué dimensiones son relevantes y cómo se protegerá la confidencialidad.",
        "Después de recopilar respuestas, los resultados deben analizarse por dimensiones y patrones. Un promedio general puede ocultar diferencias importantes entre equipos. La medición termina cuando la organización comparte resultados, define acciones, asigna responsables y vuelve a medir.",
        "La calidad de un instrumento depende de que mida de forma consistente y válida. La literatura psicométrica recomienda considerar consistencia, validez, capacidad para detectar cambios y utilidad interpretativa antes de tomar decisiones."
    ]),
    ("7. Dimensiones recomendadas", [
        "Liderazgo: confianza, claridad, apoyo y coherencia del supervisor.",
        "Comunicación: acceso a información, escucha, voz y calidad del feedback.",
        "Reconocimiento y justicia: valoración del esfuerzo y percepción de trato equitativo.",
        "Carga y recursos: presión, herramientas, tiempo y equilibrio.",
        "Seguridad psicológica: posibilidad de preguntar, disentir o reconocer errores sin miedo.",
        "Autonomía y desarrollo: capacidad de decidir y oportunidades de aprendizaje."
    ]),
    ("8. Actividad aplicada: primer diagnóstico", [
        "Elige una organización que conozcas. Observa una práctica cotidiana, por ejemplo una reunión, la asignación de tareas o la comunicación de una decisión. Describe primero lo observable sin interpretar.",
        "Después analiza la situación en los tres niveles: qué ocurre con las personas, qué ocurre en el equipo y qué condiciones organizacionales influyen. Finalmente propone una pregunta de clima que ayude a obtener evidencia adicional."
    ]),
    ("9. Recuperación activa", [
        "Sin consultar la guía, explica la diferencia entre cultura y clima.",
        "Nombra los tres niveles de análisis y ofrece un ejemplo de cada uno.",
        "Explica por qué medir clima no termina al distribuir una encuesta.",
        "Describe una situación donde el diseño del trabajo pueda afectar la motivación.",
        "Propón dos dimensiones prioritarias para una organización con alta rotación."
    ]),
]

survey_rows = [
    ("Liderazgo", "Mi supervisor comunica prioridades con claridad."),
    ("Comunicación", "Puedo expresar inquietudes y recibir una respuesta útil."),
    ("Reconocimiento", "El buen trabajo recibe reconocimiento oportuno."),
    ("Justicia", "Las decisiones que afectan al equipo se explican de manera justa."),
    ("Carga laboral", "Dispongo del tiempo y recursos necesarios para realizar mi trabajo."),
    ("Seguridad psicológica", "Puedo reconocer un error o plantear una idea sin temor."),
]

sources = [
    "Overview Organisational Behaviour CRKC 5001, Dr Patrick Reid.",
    "OpenStax, Organizational Behavior: https://openstax.org/details/books/organizational-behavior",
    "Systematic review of organizational climate measurement properties: https://pmc.ncbi.nlm.nih.gov/articles/PMC9978646/",
    "PubMed record on organizational climate: https://pubmed.ncbi.nlm.nih.gov/22856467/",
    "Organizational Climate Questionnaire: https://homepages.se.edu/cvonbergen/files/2013/01/The-Organizational-Climate-Questionnaire.pdf",
    "SHRM Organizational Climate overview: https://hr.university/shrm/organizational-climate/",
    "Practical survey-question reference: https://www.heflo.com/blog/organizational-climate-survey-questions",
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def add_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(.8)
    sec.left_margin = sec.right_margin = Inches(.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [("Title",32,INK),("Heading 1",19,INK),("Heading 2",14,MUTED)]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    header = sec.header.paragraphs[0]
    header.text = "OSCAR LEARNING OS  |  CRKC 5001"
    header.style = doc.styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = sec.footer.paragraphs[0]
    footer.text = "Módulo 1 · Lectura completa y recuperación activa"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MÓDULO 1")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string("6F8E16")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Fundamentos del\ncomportamiento organizacional")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guía completa de lectura y memoria para Oscar Vargas")
    r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CRKC 5001  ·  Definitions & Orientation  ·  Nature of Work & Management")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("6F8E16"); r.bold = True
    doc.add_page_break()

    doc.add_heading("Mapa de lectura", level=1)
    for title, _ in sections:
        p = doc.add_paragraph(style="List Number")
        p.add_run(title.split(". ",1)[1])
    doc.add_page_break()

    for idx, (title, paras) in enumerate(sections):
        doc.add_heading(title, level=1)
        for text in paras:
            doc.add_paragraph(text)
        if idx in (1, 4, 5):
            p = doc.add_paragraph()
            shade_par = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F7E7"); shade_par.append(shd)
            run = p.add_run(["Idea fuerza: conducta, contexto y resultados deben analizarse juntos.",
                             "Idea fuerza: la cultura explica la identidad; el clima muestra cómo esa identidad se vive.",
                             "Idea fuerza: medir sin actuar deteriora la confianza y reduce la participación futura."][[1,4,5].index(idx)])
            run.bold = True; run.font.color.rgb = RGBColor.from_string("526A13")
        if idx in (3, 6):
            doc.add_page_break()

    doc.add_heading("10. Instrumento inicial de práctica", level=1)
    doc.add_paragraph("Escala sugerida: 1 = totalmente en desacuerdo; 5 = totalmente de acuerdo. Este instrumento es educativo y debe validarse antes de usarse para decisiones organizacionales reales.")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(4.9)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Dimensión", "Afirmación de ejemplo"
    for cell in hdr:
        shade(cell, INK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255); run.bold = True
    for dimension, question in survey_rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = dimension, question
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_page_break()
    doc.add_heading("Fuentes y lecturas recomendadas", level=1)
    for source in sources:
        doc.add_paragraph(source, style="List Bullet")
    doc.add_paragraph("Nota de uso: esta guía contiene síntesis originales para aprendizaje. Las fuentes enlazadas deben consultarse respetando sus licencias y condiciones de uso.")
    doc.save(DOCX)

def add_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverKicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, textColor=HexColor("#718C1B"), alignment=TA_CENTER, spaceAfter=18))
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=29, leading=32, textColor=HexColor("#151820"), alignment=TA_CENTER, spaceAfter=14))
    styles.add(ParagraphStyle(name="CoverSub", parent=styles["Normal"], fontSize=12, leading=16, textColor=HexColor("#596273"), alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=HexColor("#151820"), spaceBefore=14, spaceAfter=10))
    styles.add(ParagraphStyle(name="Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.2, leading=14.5, textColor=HexColor("#232832"), spaceAfter=8))
    story = [Spacer(1,1.55*inch), Paragraph("MÓDULO 1",styles["CoverKicker"]), Paragraph("Fundamentos del<br/>comportamiento organizacional",styles["CoverTitle"]), Paragraph("Guía completa de lectura y memoria para Oscar Vargas<br/><br/>CRKC 5001 · Definitions & Orientation · Nature of Work & Management",styles["CoverSub"]), PageBreak()]
    story += [Paragraph("Mapa de lectura",styles["H1x"])]
    for title,_ in sections:
        story.append(Paragraph(title,styles["Bodyx"]))
    story.append(PageBreak())
    for idx,(title,paras) in enumerate(sections):
        story.append(Paragraph(title,styles["H1x"]))
        for text in paras:
            story.append(Paragraph(text,styles["Bodyx"]))
        if idx in (3,6):
            story.append(PageBreak())
    story.append(Paragraph("10. Instrumento inicial de práctica",styles["H1x"]))
    story.append(Paragraph("Escala sugerida: 1 = totalmente en desacuerdo; 5 = totalmente de acuerdo. Instrumento educativo que requiere validación antes de decisiones reales.",styles["Bodyx"]))
    data = [["Dimensión","Afirmación de ejemplo"]] + survey_rows
    table = Table(data, colWidths=[1.55*inch,4.65*inch], repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),HexColor("#151820")),("TEXTCOLOR",(0,0),(-1,0),white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),8.5),("LEADING",(0,0),(-1,-1),11),("GRID",(0,0),(-1,-1),.4,HexColor("#D9DEE5")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8)]))
    story.append(table)
    story.append(PageBreak())
    story.append(Paragraph("Fuentes y lecturas recomendadas",styles["H1x"]))
    for source in sources:
        story.append(Paragraph("• "+source,styles["Bodyx"]))
    def page(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica",8)
        canvas.setFillColor(HexColor("#596273"))
        canvas.drawString(.8*inch,.45*inch,"OSCAR LEARNING OS · CRKC 5001")
        canvas.drawRightString(7.7*inch,.45*inch,str(doc.page))
        canvas.restoreState()
    SimpleDocTemplate(str(PDF), pagesize=letter, rightMargin=.85*inch,leftMargin=.85*inch,topMargin=.8*inch,bottomMargin=.7*inch).build(story,onFirstPage=page,onLaterPages=page)

add_docx()
add_pdf()
print(DOCX)
print(PDF)
